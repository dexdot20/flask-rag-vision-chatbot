from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer

from canvas_service import extract_canvas_documents

_FONT_PATHS = {
    "DejaVuSans": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "DejaVuSans-Bold": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "DejaVuSansMono": "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
}


def _try_register_fonts() -> bool:
    import os

    if not all(os.path.exists(path) for path in _FONT_PATHS.values()):
        return False
    try:
        for name, path in _FONT_PATHS.items():
            pdfmetrics.registerFont(TTFont(name, path))
        return True
    except Exception:
        return False


_UNICODE_FONTS = _try_register_fonts()
_BODY_FONT = "DejaVuSans" if _UNICODE_FONTS else "Helvetica"
_BOLD_FONT = "DejaVuSans-Bold" if _UNICODE_FONTS else "Helvetica-Bold"
_MONO_FONT = "DejaVuSansMono" if _UNICODE_FONTS else "Courier"


def _escape_pdf_text(value: str) -> str:
    return str(value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_export_header(conversation: dict) -> list[str]:
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    model = str(conversation.get("model") or "").strip()
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    lines = [f"# {title}", "", f"Exported at: {exported_at}"]
    if model:
        lines.append(f"Model: {model}")
    return lines


def _iter_message_sections(messages: list[dict]) -> list[dict]:
    sections = []
    for index, message in enumerate(messages or [], start=1):
        if not isinstance(message, dict):
            continue

        role = str(message.get("role") or "message").strip() or "message"
        content = str(message.get("content") or "").strip()
        metadata = message.get("metadata") if isinstance(message.get("metadata"), dict) else {}
        details = []

        reasoning = str(metadata.get("reasoning_content") or "").strip()
        if reasoning:
            details.append(("Reasoning", reasoning))

        tool_trace = metadata.get("tool_trace") if isinstance(metadata.get("tool_trace"), list) else []
        if tool_trace:
            lines = []
            for entry in tool_trace:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                step = entry.get("step")
                summary = str(entry.get("summary") or entry.get("preview") or "").strip()
                prefix = f"Step {int(step)}" if isinstance(step, (int, float)) else "Step"
                lines.append(f"- {prefix}: {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Trace", "\n".join(lines)))

        tool_results = metadata.get("tool_results") if isinstance(metadata.get("tool_results"), list) else []
        if tool_results:
            lines = []
            for entry in tool_results:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                summary = str(entry.get("summary") or "").strip()
                lines.append(f"- {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Results", "\n".join(lines)))

        canvas_documents = extract_canvas_documents(metadata)
        if canvas_documents:
            lines = []
            for document in canvas_documents:
                title = str(document.get("title") or "Canvas").strip() or "Canvas"
                doc_content = str(document.get("content") or "").strip()
                lines.append(f"### Canvas: {title}")
                if doc_content:
                    lines.append("")
                    lines.append("```markdown")
                    lines.append(doc_content)
                    lines.append("```")
                    lines.append("")
            if lines:
                details.append(("Canvas Documents", "\n".join(lines).strip()))

        sections.append(
            {
                "title": f"## {index}. {role.title()}",
                "content": content,
                "details": details,
            }
        )
    return sections


def build_conversation_markdown_download(conversation: dict, messages: list[dict]) -> bytes:
    lines = _build_export_header(conversation)
    for section in _iter_message_sections(messages):
        lines.extend(["", section["title"], ""])
        lines.append(section["content"] or "_(empty)_")
        for label, value in section["details"]:
            lines.extend(["", f"### {label}", "", value])
    lines.append("")
    return "\n".join(lines).encode("utf-8")


def build_conversation_docx_download(conversation: dict, messages: list[dict]) -> bytes:
    document = Document()
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    document.add_heading(title, level=0)
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    meta_parts = [f"Exported at: {exported_at}"]
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    document.add_paragraph(" | ".join(meta_parts))

    for section in _iter_message_sections(messages):
        document.add_heading(section["title"].replace("## ", ""), level=1)
        document.add_paragraph(section["content"] or "(empty)")
        for label, value in section["details"]:
            document.add_heading(label, level=2)
            for block in str(value).split("\n\n"):
                stripped = block.strip()
                if stripped:
                    document.add_paragraph(stripped)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_conversation_pdf_download(conversation: dict, messages: list[dict]) -> bytes:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ConversationTitle", parent=styles["Title"], fontName=_BOLD_FONT)
    heading_style = ParagraphStyle(
        "ConversationHeading",
        parent=styles["Heading2"],
        fontName=_BOLD_FONT,
        textColor=colors.HexColor("#1f2a44"),
        spaceAfter=8,
        spaceBefore=12,
    )
    subheading_style = ParagraphStyle(
        "ConversationSubheading",
        parent=styles["Heading3"],
        fontName=_BOLD_FONT,
        textColor=colors.HexColor("#33415f"),
        spaceAfter=6,
        spaceBefore=8,
    )
    body_style = ParagraphStyle(
        "ConversationBody",
        parent=styles["BodyText"],
        fontName=_BODY_FONT,
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "ConversationCode",
        parent=styles["Code"],
        fontName=_MONO_FONT,
        fontSize=8.5,
        leading=11,
        leftIndent=10,
        rightIndent=10,
        backColor=colors.HexColor("#f3f5f9"),
        borderPadding=8,
    )

    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    story = [Paragraph(_escape_pdf_text(title), title_style), Spacer(1, 6)]
    meta_parts = [f"Exported at: {exported_at}"]
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    story.append(Paragraph(_escape_pdf_text(" | ".join(meta_parts)), body_style))

    for section in _iter_message_sections(messages):
        story.append(Spacer(1, 6))
        story.append(Paragraph(_escape_pdf_text(section["title"].replace("## ", "")), heading_style))
        story.append(Paragraph(_escape_pdf_text(section["content"] or "(empty)").replace("\n", "<br/>"), body_style))
        for label, value in section["details"]:
            story.append(Paragraph(_escape_pdf_text(label), subheading_style))
            stripped_value = str(value or "").strip()
            if "```" in stripped_value:
                code_lines = []
                in_code_block = False
                for line in stripped_value.splitlines():
                    if line.strip().startswith("```"):
                        if in_code_block and code_lines:
                            story.append(Preformatted("\n".join(code_lines), code_style))
                            story.append(Spacer(1, 4))
                            code_lines = []
                        in_code_block = not in_code_block
                        continue
                    if in_code_block:
                        code_lines.append(line)
                    else:
                        story.append(Paragraph(_escape_pdf_text(line or " "), body_style))
                if code_lines:
                    story.append(Preformatted("\n".join(code_lines), code_style))
            elif stripped_value.startswith("- "):
                items = [line[2:].strip() for line in stripped_value.splitlines() if line.strip().startswith("- ")]
                if items:
                    story.append(
                        ListFlowable(
                            [ListItem(Paragraph(_escape_pdf_text(item), body_style)) for item in items],
                            bulletType="bullet",
                            leftIndent=18,
                        )
                    )
                    story.append(Spacer(1, 4))
            else:
                story.append(Paragraph(_escape_pdf_text(stripped_value).replace("\n", "<br/>"), body_style))

    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    doc.build(story)
    return output.getvalue()