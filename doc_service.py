from __future__ import annotations

import csv
import io
import os

import pdfplumber
from docx import Document

from config import (
    DOCUMENT_ALLOWED_MIME_TYPES,
    DOCUMENT_MAX_BYTES,
    DOCUMENT_MAX_TEXT_CHARS,
)

MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_PDF = "application/pdf"
MIME_PLAIN = "text/plain"
MIME_CSV = "text/csv"
MIME_MARKDOWN = "text/markdown"

_EXTENSION_TO_MIME: dict[str, str] = {
    ".docx": MIME_DOCX,
    ".pdf": MIME_PDF,
    ".txt": MIME_PLAIN,
    ".csv": MIME_CSV,
    ".md": MIME_MARKDOWN,
    ".py": MIME_PLAIN,
    ".js": MIME_PLAIN,
    ".ts": MIME_PLAIN,
    ".tsx": MIME_PLAIN,
    ".jsx": MIME_PLAIN,
    ".json": MIME_PLAIN,
    ".html": MIME_PLAIN,
    ".css": MIME_PLAIN,
    ".scss": MIME_PLAIN,
    ".sh": MIME_PLAIN,
    ".sql": MIME_PLAIN,
    ".yaml": MIME_PLAIN,
    ".yml": MIME_PLAIN,
}

_CODE_LANGUAGE_BY_EXTENSION: dict[str, str] = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".tsx": "tsx",
    ".jsx": "jsx",
    ".json": "json",
    ".html": "html",
    ".css": "css",
    ".scss": "scss",
    ".sh": "bash",
    ".sql": "sql",
    ".yaml": "yaml",
    ".yml": "yaml",
}


def guess_document_mime_type(filename: str, declared_mime: str) -> str:
    declared = (declared_mime or "").strip().lower()
    if declared in DOCUMENT_ALLOWED_MIME_TYPES:
        return declared
    ext = os.path.splitext(filename or "")[-1].lower()
    return _EXTENSION_TO_MIME.get(ext, declared)


def read_uploaded_document(uploaded_file) -> tuple[str, str, bytes]:
    filename = os.path.basename((uploaded_file.filename or "").strip())
    declared_mime = (uploaded_file.mimetype or "").lower().strip()
    mime_type = guess_document_mime_type(filename, declared_mime)
    if mime_type not in DOCUMENT_ALLOWED_MIME_TYPES:
        raise ValueError("Unsupported document type. Upload DOCX, PDF, TXT, CSV or MD.")
    doc_bytes = uploaded_file.read()
    if not doc_bytes:
        raise ValueError("Uploaded document is empty.")
    if len(doc_bytes) > DOCUMENT_MAX_BYTES:
        raise ValueError(f"Document is too large. Upload a maximum of {DOCUMENT_MAX_BYTES // (1024 * 1024)} MB.")
    return filename, mime_type, doc_bytes


def _extract_text_from_docx(doc_bytes: bytes) -> str:
    document = Document(io.BytesIO(doc_bytes))
    paragraphs = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_text_from_pdf(doc_bytes: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(doc_bytes)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def _extract_text_plain(doc_bytes: bytes) -> str:
    return doc_bytes.decode("utf-8", errors="replace").strip()


def _extract_text_csv(doc_bytes: bytes) -> str:
    text = doc_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows: list[str] = []
    for row in reader:
        stripped = [cell.strip() for cell in row]
        if any(stripped):
            rows.append(" | ".join(stripped))
    return "\n".join(rows)


def extract_document_text(doc_bytes: bytes, mime_type: str) -> str:
    mime = (mime_type or "").strip().lower()
    if mime == MIME_DOCX:
        return _extract_text_from_docx(doc_bytes)
    if mime == MIME_PDF:
        return _extract_text_from_pdf(doc_bytes)
    if mime == MIME_CSV:
        return _extract_text_csv(doc_bytes)
    if mime in (MIME_PLAIN, MIME_MARKDOWN):
        return _extract_text_plain(doc_bytes)
    raise ValueError(f"No text extractor for MIME type: {mime}")


def infer_canvas_language(filename: str) -> str | None:
    ext = os.path.splitext(filename or "")[-1].lower()
    return _CODE_LANGUAGE_BY_EXTENSION.get(ext)


def infer_canvas_format(filename: str) -> str:
    return "code" if infer_canvas_language(filename) else "markdown"


def build_canvas_markdown(filename: str, text: str) -> str:
    name = os.path.basename(filename or "document")
    if infer_canvas_format(name) == "code":
        return text.rstrip("\n")
    if os.path.splitext(name)[-1].lower() == ".md":
        return text
    return f"# {name}\n\n{text}"


def build_document_context_block(filename: str, text: str) -> tuple[str, bool]:
    truncated = len(text) > DOCUMENT_MAX_TEXT_CHARS
    clipped = text[:DOCUMENT_MAX_TEXT_CHARS] if truncated else text
    header = f"[Uploaded document: {os.path.basename(filename or 'document')}]"
    if truncated:
        header += " (truncated to first 50,000 characters)"
    return f"{header}\n{clipped}", truncated
